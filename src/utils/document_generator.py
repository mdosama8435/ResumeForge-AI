import io
from typing import Dict, Any
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

class DocumentGenerator:
    @staticmethod
    def generate_pdf(opt_data: Dict[str, Any]) -> bytes:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        styles = getSampleStyleSheet()
        story = []
        
        title_style = styles['Heading1']
        heading_style = styles['Heading2']
        body_style = styles['Normal']
        
        story.append(Paragraph("Optimized Resume", title_style))
        story.append(Spacer(1, 12))
        
        if summary := opt_data.get("summary"):
            story.append(Paragraph("Professional Summary", heading_style))
            story.append(Paragraph(summary, body_style))
            story.append(Spacer(1, 12))
            
        if skills := opt_data.get("skills"):
            story.append(Paragraph("Skills", heading_style))
            story.append(Paragraph(", ".join(skills), body_style))
            story.append(Spacer(1, 12))
            
        if experience := opt_data.get("experience"):
            story.append(Paragraph("Experience", heading_style))
            for exp in experience:
                story.append(Paragraph(f"• {exp}", body_style))
            story.append(Spacer(1, 12))
            
        if projects := opt_data.get("projects"):
            story.append(Paragraph("Projects", heading_style))
            for proj in projects:
                story.append(Paragraph(f"• {proj}", body_style))
                
        doc.build(story)
        return buffer.getvalue()

    @staticmethod
    def generate_docx(opt_data: Dict[str, Any]) -> bytes:
        doc = Document()
        doc.add_heading('Optimized Resume', 0)
        
        if summary := opt_data.get("summary"):
            doc.add_heading('Professional Summary', level=1)
            doc.add_paragraph(summary)
            
        if skills := opt_data.get("skills"):
            doc.add_heading('Skills', level=1)
            doc.add_paragraph(", ".join(skills))
            
        if experience := opt_data.get("experience"):
            doc.add_heading('Experience', level=1)
            for exp in experience:
                doc.add_paragraph(exp, style='List Bullet')
                
        if projects := opt_data.get("projects"):
            doc.add_heading('Projects', level=1)
            for proj in projects:
                doc.add_paragraph(proj, style='List Bullet')
                
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
